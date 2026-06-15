from __future__ import annotations

import hashlib
import hmac
import io
import os
import secrets
import sqlite3
from contextlib import contextmanager
from datetime import datetime
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


APP_DIR = Path(__file__).resolve().parent
DB_PATH = APP_DIR / "dentalcrm.db"
UPLOAD_DIR = APP_DIR / "uploads"
STL_DIR = UPLOAD_DIR / "STL"
PHOTO_DIR = UPLOAD_DIR / "Photos"
INVOICE_DIR = UPLOAD_DIR / "Invoices"
TEMPLATE_DIR = APP_DIR / "Invoice Templates"

ORDER_STATUSES = [
    "New",
    "STL Received",
    "In Production",
    "Quality Check",
    "Ready",
    "Dispatched",
    "Completed",
    "Cancelled",
]
PAYMENT_STATUSES = ["Unpaid", "Partial", "Paid", "Refunded"]


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def money(value: float) -> str:
    return f"INR {value:,.2f}"


def hash_password(password: str, salt: str | None = None) -> str:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac(
        "sha256", password.encode("utf-8"), bytes.fromhex(salt), 200_000
    ).hex()
    return f"{salt}${digest}"


def verify_password(password: str, stored: str) -> bool:
    try:
        salt, expected = stored.split("$", 1)
    except ValueError:
        return False
    actual = hash_password(password, salt).split("$", 1)[1]
    return hmac.compare_digest(actual, expected)


@contextmanager
def db():
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    connection.execute("PRAGMA foreign_keys = ON")
    try:
        yield connection
        connection.commit()
    finally:
        connection.close()


def initialize_app() -> None:
    for folder in (STL_DIR, PHOTO_DIR, INVOICE_DIR, TEMPLATE_DIR):
        folder.mkdir(parents=True, exist_ok=True)
    create_invoice_templates()

    with db() as connection:
        connection.executescript(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                full_name TEXT NOT NULL,
                email TEXT NOT NULL UNIQUE COLLATE NOCASE,
                phone TEXT DEFAULT '',
                clinic_name TEXT DEFAULT '',
                address TEXT DEFAULT '',
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL CHECK(role IN ('admin', 'doctor')),
                approved INTEGER NOT NULL DEFAULT 0,
                created_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS orders (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                order_number TEXT NOT NULL UNIQUE,
                doctor_id INTEGER NOT NULL,
                patient_name TEXT NOT NULL,
                patient_code TEXT DEFAULT '',
                service_type TEXT NOT NULL,
                tooth_details TEXT DEFAULT '',
                shade TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                status TEXT NOT NULL DEFAULT 'New',
                amount REAL NOT NULL DEFAULT 0,
                payment_status TEXT NOT NULL DEFAULT 'Unpaid',
                amount_paid REAL NOT NULL DEFAULT 0,
                due_date TEXT DEFAULT '',
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY (doctor_id) REFERENCES users(id)
            );

            CREATE TABLE IF NOT EXISTS uploads (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                order_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                category TEXT NOT NULL CHECK(category IN ('STL', 'Photo')),
                original_name TEXT NOT NULL,
                stored_path TEXT NOT NULL,
                uploaded_at TEXT NOT NULL,
                FOREIGN KEY (order_id) REFERENCES orders(id) ON DELETE CASCADE,
                FOREIGN KEY (user_id) REFERENCES users(id)
            );

            CREATE TABLE IF NOT EXISTS invoices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_number TEXT NOT NULL UNIQUE,
                order_id INTEGER NOT NULL UNIQUE,
                subtotal REAL NOT NULL,
                tax_rate REAL NOT NULL,
                total REAL NOT NULL,
                pdf_path TEXT NOT NULL,
                docx_path TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (order_id) REFERENCES orders(id)
            );

            CREATE TABLE IF NOT EXISTS payments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                order_id INTEGER NOT NULL,
                amount REAL NOT NULL,
                method TEXT NOT NULL,
                reference TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                paid_at TEXT NOT NULL,
                recorded_by INTEGER NOT NULL,
                FOREIGN KEY (order_id) REFERENCES orders(id),
                FOREIGN KEY (recorded_by) REFERENCES users(id)
            );
            """
        )
        admin = connection.execute(
            "SELECT id FROM users WHERE email = ?", ("admin@dentalcrm.local",)
        ).fetchone()
        if not admin:
            connection.execute(
                """
                INSERT INTO users
                    (full_name, email, phone, clinic_name, address, password_hash,
                     role, approved, created_at)
                VALUES (?, ?, ?, ?, ?, ?, 'admin', 1, ?)
                """,
                (
                    "Dental CRM Admin",
                    "admin@dentalcrm.local",
                    "",
                    "Dental Laboratory",
                    "",
                    hash_password("Admin@123"),
                    now(),
                ),
            )


def set_docx_font(run, size: float, color: str = "123B5D", bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def create_invoice_templates() -> None:
    pdf_path = TEMPLATE_DIR / "Invoice.pdf"
    docx_path = TEMPLATE_DIR / "Invoice.docx"
    rows = [
        ["Description", "Amount"],
        ["{{SERVICE_TYPE}} - {{TOOTH_DETAILS}}", "{{SUBTOTAL}}"],
        ["Tax ({{TAX_RATE}}%)", "{{TAX_AMOUNT}}"],
        ["Total", "{{TOTAL}}"],
        ["Paid", "{{AMOUNT_PAID}}"],
        ["Balance", "{{BALANCE}}"],
    ]

    if not pdf_path.exists():
        buffer = io.BytesIO()
        document = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
        )
        styles = getSampleStyleSheet()
        story = [
            Paragraph("DENTAL LABORATORY", styles["Title"]),
            Paragraph("INVOICE TEMPLATE", styles["Heading2"]),
            Spacer(1, 8),
            Table(
                [
                    ["Invoice No.", "{{INVOICE_NUMBER}}", "Date", "{{INVOICE_DATE}}"],
                    ["Order No.", "{{ORDER_NUMBER}}", "Due Date", "{{DUE_DATE}}"],
                    ["Doctor", "{{DOCTOR_NAME}}", "Clinic", "{{CLINIC_NAME}}"],
                    ["Patient", "{{PATIENT_NAME}}", "Service", "{{SERVICE_TYPE}}"],
                ],
                colWidths=[28 * mm, 55 * mm, 28 * mm, 55 * mm],
                style=TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")),
                        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#F2F4F7")),
                        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
                        ("PADDING", (0, 0), (-1, -1), 7),
                    ]
                ),
            ),
            Spacer(1, 14),
            Table(
                rows,
                colWidths=[120 * mm, 45 * mm],
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#123B5D")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#E0F2FE")),
                        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                        ("PADDING", (0, 0), (-1, -1), 8),
                    ]
                ),
            ),
            Spacer(1, 20),
            Paragraph("Thank you for your business.", styles["Normal"]),
        ]
        document.build(story)
        pdf_path.write_bytes(buffer.getvalue())

    if not docx_path.exists():
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)

        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title.paragraph_format.space_after = Pt(4)
        set_docx_font(title.add_run("DENTAL LABORATORY"), 24, bold=True)
        subtitle = document.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(16)
        set_docx_font(subtitle.add_run("Invoice Template"), 14, "64748B", bold=True)

        info = document.add_table(rows=4, cols=4)
        info.style = "Table Grid"
        info.autofit = False
        widths = [Inches(1.0), Inches(2.25), Inches(1.0), Inches(2.25)]
        info_data = [
            ("Invoice No.", "{{INVOICE_NUMBER}}", "Date", "{{INVOICE_DATE}}"),
            ("Order No.", "{{ORDER_NUMBER}}", "Due Date", "{{DUE_DATE}}"),
            ("Doctor", "{{DOCTOR_NAME}}", "Clinic", "{{CLINIC_NAME}}"),
            ("Patient", "{{PATIENT_NAME}}", "Service", "{{SERVICE_TYPE}}"),
        ]
        for table_row, values in zip(info.rows, info_data):
            for cell, value, width in zip(table_row.cells, values, widths):
                cell.width = width
                cell.text = value
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
        document.add_paragraph()

        totals = document.add_table(rows=len(rows), cols=2)
        totals.style = "Table Grid"
        totals.autofit = False
        for table_row, values in zip(totals.rows, rows):
            table_row.cells[0].width = Inches(5)
            table_row.cells[1].width = Inches(1.5)
            table_row.cells[0].text, table_row.cells[1].text = values
            table_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cell in table_row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
        for cell in totals.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                set_docx_font(run, 11, "FFFFFF", bold=True)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "123B5D")
            cell._tc.get_or_add_tcPr().append(shading)

        closing = document.add_paragraph()
        closing.paragraph_format.space_before = Pt(18)
        closing.add_run("Thank you for your business.")
        document.save(docx_path)


def query(sql: str, params: tuple = ()) -> list[sqlite3.Row]:
    with db() as connection:
        return connection.execute(sql, params).fetchall()


def query_one(sql: str, params: tuple = ()) -> sqlite3.Row | None:
    rows = query(sql, params)
    return rows[0] if rows else None


def execute(sql: str, params: tuple = ()) -> int:
    with db() as connection:
        cursor = connection.execute(sql, params)
        return cursor.lastrowid


def make_order_number() -> str:
    prefix = datetime.now().strftime("ORD-%Y%m%d")
    row = query_one(
        "SELECT COUNT(*) AS count FROM orders WHERE order_number LIKE ?",
        (f"{prefix}-%",),
    )
    return f"{prefix}-{row['count'] + 1:03d}"


def make_invoice_number() -> str:
    prefix = datetime.now().strftime("INV-%Y%m")
    row = query_one(
        "SELECT COUNT(*) AS count FROM invoices WHERE invoice_number LIKE ?",
        (f"{prefix}-%",),
    )
    return f"{prefix}-{row['count'] + 1:04d}"


def safe_filename(name: str) -> str:
    clean = "".join(c for c in Path(name).name if c.isalnum() or c in "._- ")
    return clean.strip() or "upload"


def save_upload(order_id: int, user_id: int, category: str, uploaded_file) -> None:
    destination_dir = STL_DIR if category == "STL" else PHOTO_DIR
    order = query_one("SELECT order_number FROM orders WHERE id = ?", (order_id,))
    folder = destination_dir / order["order_number"]
    folder.mkdir(parents=True, exist_ok=True)
    filename = f"{datetime.now():%Y%m%d%H%M%S%f}_{safe_filename(uploaded_file.name)}"
    path = folder / filename
    path.write_bytes(uploaded_file.getbuffer())
    execute(
        """
        INSERT INTO uploads
            (order_id, user_id, category, original_name, stored_path, uploaded_at)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (order_id, user_id, category, uploaded_file.name, str(path), now()),
    )
    if category == "STL":
        execute(
            "UPDATE orders SET status = 'STL Received', updated_at = ? WHERE id = ?",
            (now(), order_id),
        )


def invoice_context(order_id: int) -> sqlite3.Row:
    return query_one(
        """
        SELECT o.*, u.full_name AS doctor_name, u.clinic_name, u.email, u.phone,
               u.address
        FROM orders o
        JOIN users u ON u.id = o.doctor_id
        WHERE o.id = ?
        """,
        (order_id,),
    )


def generate_invoice_files(
    order_id: int,
    tax_rate: float,
    invoice_number: str | None = None,
    invoice_date=None,
) -> sqlite3.Row:
    order = invoice_context(order_id)
    invoice_number = (invoice_number or "").strip() or make_invoice_number()
    invoice_date = invoice_date or datetime.now().date()
    invoice_date_text = invoice_date.strftime("%d %b %Y")
    created_at = f"{invoice_date.isoformat()} {datetime.now():%H:%M:%S}"
    subtotal = float(order["amount"])
    tax = subtotal * tax_rate / 100
    total = subtotal + tax
    invoice_filename = safe_filename(invoice_number)
    pdf_path = INVOICE_DIR / f"{invoice_filename}.pdf"
    docx_path = INVOICE_DIR / f"{invoice_filename}.docx"

    pdf_buffer = io.BytesIO()
    document = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    styles = getSampleStyleSheet()
    story = [
        Paragraph("DENTAL LABORATORY", styles["Title"]),
        Paragraph("TAX INVOICE", styles["Heading2"]),
        Spacer(1, 8),
        Table(
            [
                ["Invoice No.", invoice_number, "Date", invoice_date_text],
                ["Order No.", order["order_number"], "Due Date", order["due_date"] or "-"],
                ["Doctor", order["doctor_name"], "Clinic", order["clinic_name"] or "-"],
                ["Patient", order["patient_name"], "Service", order["service_type"]],
            ],
            colWidths=[28 * mm, 55 * mm, 28 * mm, 55 * mm],
        ),
        Spacer(1, 14),
        Table(
            [
                ["Description", "Amount"],
                [
                    f"{order['service_type']} - {order['tooth_details'] or 'Dental case'}",
                    money(subtotal),
                ],
                [f"Tax ({tax_rate:.2f}%)", money(tax)],
                ["Total", money(total)],
                ["Paid", money(float(order["amount_paid"]))],
                ["Balance", money(max(total - float(order["amount_paid"]), 0))],
            ],
            colWidths=[120 * mm, 45 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#123B5D")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                    ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#E0F2FE")),
                    ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                    ("PADDING", (0, 0), (-1, -1), 8),
                ]
            ),
        ),
        Spacer(1, 20),
        Paragraph("Thank you for your business.", styles["Normal"]),
    ]
    document.build(story)
    pdf_path.write_bytes(pdf_buffer.getvalue())

    docx = Document()
    docx.add_heading("DENTAL LABORATORY", 0)
    docx.add_heading("Tax Invoice", level=1)
    info = docx.add_table(rows=4, cols=4)
    info.style = "Table Grid"
    info_data = [
        ("Invoice No.", invoice_number, "Date", invoice_date_text),
        ("Order No.", order["order_number"], "Due Date", order["due_date"] or "-"),
        ("Doctor", order["doctor_name"], "Clinic", order["clinic_name"] or "-"),
        ("Patient", order["patient_name"], "Service", order["service_type"]),
    ]
    for row, values in zip(info.rows, info_data):
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
    docx.add_paragraph()
    totals = docx.add_table(rows=6, cols=2)
    totals.style = "Table Grid"
    total_data = [
        ("Description", "Amount"),
        (f"{order['service_type']} - {order['tooth_details'] or 'Dental case'}", money(subtotal)),
        (f"Tax ({tax_rate:.2f}%)", money(tax)),
        ("Total", money(total)),
        ("Paid", money(float(order["amount_paid"]))),
        ("Balance", money(max(total - float(order["amount_paid"]), 0))),
    ]
    for row, values in zip(totals.rows, total_data):
        row.cells[0].text, row.cells[1].text = values
    docx.add_paragraph("Thank you for your business.")
    docx.save(docx_path)

    execute(
        """
        INSERT INTO invoices
            (invoice_number, order_id, subtotal, tax_rate, total, pdf_path,
             docx_path, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            invoice_number,
            order_id,
            subtotal,
            tax_rate,
            total,
            str(pdf_path),
            str(docx_path),
            created_at,
        ),
    )
    return query_one("SELECT * FROM invoices WHERE order_id = ?", (order_id,))


def doctor_options(approved_only: bool = True) -> dict[str, int]:
    approval_clause = "AND approved = 1" if approved_only else ""
    doctors = query(
        f"""
        SELECT id, full_name, clinic_name, email
        FROM users
        WHERE role = 'doctor' {approval_clause}
        ORDER BY full_name
        """
    )
    return {
        (
            f"{doctor['full_name']} | {doctor['clinic_name'] or 'No clinic'} | "
            f"{doctor['email']}"
        ): doctor["id"]
        for doctor in doctors
    }


def doctor_lookup(
    label: str,
    key: str,
    approved_only: bool = True,
    allow_all: bool = False,
) -> int | None:
    options = doctor_options(approved_only)
    if allow_all:
        options = {"All doctors": None, **options}
    if not options:
        st.info("No matching doctor accounts are available.")
        return None
    selected = st.selectbox(
        label,
        options,
        key=key,
        help="Click the field and type a doctor name, clinic, or email to search.",
    )
    return options[selected]


def create_order_form(doctor_id: int | None = None) -> None:
    doctors = doctor_options()
    if doctor_id is None and not doctors:
        st.info("Approve at least one doctor before creating an order.")
        return

    with st.form("create_order", clear_on_submit=True):
        if doctor_id is None:
            doctor_label = st.selectbox(
                "Search and select doctor *",
                doctors,
                help="Type a doctor name, clinic, or email to search.",
            )
            selected_doctor = doctors[doctor_label]
        else:
            selected_doctor = doctor_id

        col1, col2 = st.columns(2)
        patient_name = col1.text_input("Patient name *")
        patient_code = col2.text_input("Patient ID / code")
        service_type = col1.selectbox(
            "Service *",
            [
                "Crown",
                "Bridge",
                "Implant",
                "Denture",
                "Aligner",
                "Night Guard",
                "Digital Design",
                "Other",
            ],
        )
        tooth_details = col2.text_input("Tooth / arch details")
        shade = col1.text_input("Shade")
        due_date = col2.date_input("Due date", value=None)
        amount = col1.number_input("Order amount (INR)", min_value=0.0, step=500.0)
        notes = st.text_area("Clinical notes")
        submitted = st.form_submit_button("Create order", type="primary")
        if submitted:
            if not patient_name.strip():
                st.error("Patient name is required.")
            else:
                timestamp = now()
                execute(
                    """
                    INSERT INTO orders
                        (order_number, doctor_id, patient_name, patient_code,
                         service_type, tooth_details, shade, notes, amount,
                         due_date, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        make_order_number(),
                        selected_doctor,
                        patient_name.strip(),
                        patient_code.strip(),
                        service_type,
                        tooth_details.strip(),
                        shade.strip(),
                        notes.strip(),
                        amount,
                        due_date.isoformat() if due_date else "",
                        timestamp,
                        timestamp,
                    ),
                )
                st.success("Order created.")
                st.rerun()


def order_selector(
    doctor_id: int | None = None,
    key: str = "order",
    uninvoiced_only: bool = False,
) -> tuple[int | None, dict]:
    where = "WHERE o.doctor_id = ?" if doctor_id else ""
    params = (doctor_id,) if doctor_id else ()
    invoice_clause = (
        "AND NOT EXISTS (SELECT 1 FROM invoices i WHERE i.order_id = o.id)"
        if uninvoiced_only and where
        else "WHERE NOT EXISTS (SELECT 1 FROM invoices i WHERE i.order_id = o.id)"
        if uninvoiced_only
        else ""
    )
    orders = query(
        f"""
        SELECT o.id, o.order_number, o.patient_name, o.status,
               u.full_name AS doctor_name, u.clinic_name
        FROM orders o
        JOIN users u ON u.id = o.doctor_id
        {where} {invoice_clause}
        ORDER BY o.created_at DESC
        """,
        params,
    )
    labels = {
        (
            f"{o['order_number']} | {o['doctor_name']} | "
            f"{o['patient_name']} | {o['status']}"
        ): o["id"]
        for o in orders
    }
    if not labels:
        st.info("No orders available.")
        return None, {}
    selected = st.selectbox(
        "Search and select order",
        labels,
        key=key,
        help="Type an order number, doctor name, or patient name to search.",
    )
    return labels[selected], labels


def upload_panel(user: sqlite3.Row, category: str) -> None:
    st.subheader(f"Upload {category}")
    selected_doctor = (
        doctor_lookup(
            "Search doctor",
            f"{category}_doctor_lookup",
            allow_all=True,
        )
        if user["role"] == "admin"
        else user["id"]
    )
    order_id, _ = order_selector(
        selected_doctor,
        key=f"{category}_order",
    )
    if not order_id:
        return
    file_types = ["stl"] if category == "STL" else ["jpg", "jpeg", "png", "webp"]
    files = st.file_uploader(
        f"Choose {category} files",
        type=file_types,
        accept_multiple_files=True,
        key=f"{category}_files",
    )
    if st.button(f"Upload {category}", type="primary", disabled=not files):
        for uploaded_file in files:
            save_upload(order_id, user["id"], category, uploaded_file)
        st.success(f"{len(files)} file(s) uploaded.")
        st.rerun()


def orders_table(doctor_id: int | None = None) -> pd.DataFrame:
    where = "WHERE o.doctor_id = ?" if doctor_id else ""
    params = (doctor_id,) if doctor_id else ()
    rows = query(
        f"""
        SELECT o.order_number AS "Order", u.full_name AS "Doctor",
               o.patient_name AS "Patient", o.service_type AS "Service",
               o.status AS "Status", o.payment_status AS "Payment",
               o.amount AS "Amount", o.amount_paid AS "Paid",
               o.due_date AS "Due", o.created_at AS "Created"
        FROM orders o
        JOIN users u ON u.id = o.doctor_id
        {where}
        ORDER BY o.created_at DESC
        """,
        params,
    )
    return pd.DataFrame([dict(row) for row in rows])


def login_page() -> None:
    st.markdown('<div class="brand">Dental CRM</div>', unsafe_allow_html=True)
    st.caption("Laboratory orders, files, invoices, and payments in one place.")
    login_tab, signup_tab = st.tabs(["Login", "Doctor signup"])

    with login_tab:
        with st.form("login"):
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Login", type="primary", use_container_width=True)
            if submitted:
                user = query_one("SELECT * FROM users WHERE email = ?", (email.strip(),))
                if not user or not verify_password(password, user["password_hash"]):
                    st.error("Invalid email or password.")
                elif user["role"] == "doctor" and not user["approved"]:
                    st.warning("Your account is awaiting admin approval.")
                else:
                    st.session_state.user_id = user["id"]
                    st.rerun()

    with signup_tab:
        with st.form("signup", clear_on_submit=True):
            name = st.text_input("Full name *")
            clinic = st.text_input("Clinic name *")
            email = st.text_input("Email *")
            phone = st.text_input("Phone")
            address = st.text_area("Address")
            password = st.text_input("Password *", type="password")
            confirm = st.text_input("Confirm password *", type="password")
            submitted = st.form_submit_button("Create account", type="primary")
            if submitted:
                if not all([name.strip(), clinic.strip(), email.strip(), password]):
                    st.error("Complete all required fields.")
                elif password != confirm:
                    st.error("Passwords do not match.")
                elif len(password) < 8:
                    st.error("Password must be at least 8 characters.")
                else:
                    try:
                        execute(
                            """
                            INSERT INTO users
                                (full_name, email, phone, clinic_name, address,
                                 password_hash, role, approved, created_at)
                            VALUES (?, ?, ?, ?, ?, ?, 'doctor', 0, ?)
                            """,
                            (
                                name.strip(),
                                email.strip().lower(),
                                phone.strip(),
                                clinic.strip(),
                                address.strip(),
                                hash_password(password),
                                now(),
                            ),
                        )
                        st.success("Account created. An admin must approve it before login.")
                    except sqlite3.IntegrityError:
                        st.error("An account with this email already exists.")


def admin_dashboard() -> None:
    pending = query_one(
        "SELECT COUNT(*) AS count FROM users WHERE role = 'doctor' AND approved = 0"
    )
    frame = orders_table()

    with st.expander("Dashboard filters", expanded=True):
        first, second, third = st.columns(3)
        doctor_names = ["All doctors"]
        if not frame.empty:
            doctor_names += sorted(frame["Doctor"].dropna().unique().tolist())
        selected_doctor = first.selectbox(
            "Doctor",
            doctor_names,
            help="Click and type a doctor name to search.",
        )
        selected_status = second.multiselect("Order status", ORDER_STATUSES)
        selected_payment = third.multiselect("Payment status", PAYMENT_STATUSES)

        fourth, fifth, sixth = st.columns([2, 1, 1])
        search = fourth.text_input(
            "Search orders",
            placeholder="Order number, doctor, patient, or service",
        ).strip()
        start_date = fifth.date_input("Created from", value=None)
        end_date = sixth.date_input("Created to", value=None)

    filtered = frame.copy()
    if not filtered.empty:
        if selected_doctor != "All doctors":
            filtered = filtered[filtered["Doctor"] == selected_doctor]
        if selected_status:
            filtered = filtered[filtered["Status"].isin(selected_status)]
        if selected_payment:
            filtered = filtered[filtered["Payment"].isin(selected_payment)]
        if search:
            searchable = filtered[["Order", "Doctor", "Patient", "Service"]].fillna("")
            mask = searchable.apply(
                lambda column: column.astype(str).str.contains(
                    search, case=False, regex=False
                )
            ).any(axis=1)
            filtered = filtered[mask]
        created = pd.to_datetime(filtered["Created"], errors="coerce")
        if start_date:
            filtered = filtered[created.dt.date >= start_date]
            created = pd.to_datetime(filtered["Created"], errors="coerce")
        if end_date:
            filtered = filtered[created.dt.date <= end_date]

    active = (
        0
        if filtered.empty
        else (~filtered["Status"].isin(["Completed", "Cancelled"])).sum()
    )
    outstanding = (
        0
        if filtered.empty
        else (filtered["Amount"] - filtered["Paid"]).clip(lower=0).sum()
    )
    cols = st.columns(4)
    cols[0].metric("Filtered orders", len(filtered))
    cols[1].metric("Active orders", int(active))
    cols[2].metric("Pending doctors", pending["count"] or 0)
    cols[3].metric("Outstanding", money(float(outstanding)))

    if filtered.empty:
        st.info("No orders match the selected filters.")
    else:
        status_counts = (
            filtered["Status"].value_counts().rename_axis("Status").reset_index(name="Orders")
        )
        left, right = st.columns([1, 2])
        left.bar_chart(status_counts, x="Status", y="Orders")
        right.dataframe(filtered, use_container_width=True, hide_index=True)


def doctor_dashboard(user_id: int) -> None:
    frame = orders_table(user_id)
    total = len(frame)
    active = 0 if frame.empty else (~frame["Status"].isin(["Completed", "Cancelled"])).sum()
    balance = 0 if frame.empty else (frame["Amount"] - frame["Paid"]).clip(lower=0).sum()
    cols = st.columns(3)
    cols[0].metric("My orders", total)
    cols[1].metric("Active cases", active)
    cols[2].metric("Balance", money(float(balance)))
    if frame.empty:
        st.info("Create your first order to begin.")
    else:
        st.dataframe(frame, use_container_width=True, hide_index=True)


def manage_doctors() -> None:
    st.subheader("Doctors")
    add_tab, directory_tab = st.tabs(["Add doctor", "Doctor directory"])

    with add_tab:
        with st.form("admin_add_doctor", clear_on_submit=True):
            col1, col2 = st.columns(2)
            name = col1.text_input("Full name *")
            clinic = col2.text_input("Clinic name *")
            email = col1.text_input("Email *")
            phone = col2.text_input("Phone")
            address = st.text_area("Address")
            password = col1.text_input("Temporary password *", type="password")
            approved = col2.checkbox("Approve immediately", value=True)
            submitted = st.form_submit_button("Add doctor", type="primary")
            if submitted:
                if not all([name.strip(), clinic.strip(), email.strip(), password]):
                    st.error("Complete all required fields.")
                elif len(password) < 8:
                    st.error("Temporary password must be at least 8 characters.")
                else:
                    try:
                        execute(
                            """
                            INSERT INTO users
                                (full_name, email, phone, clinic_name, address,
                                 password_hash, role, approved, created_at)
                            VALUES (?, ?, ?, ?, ?, ?, 'doctor', ?, ?)
                            """,
                            (
                                name.strip(),
                                email.strip().lower(),
                                phone.strip(),
                                clinic.strip(),
                                address.strip(),
                                hash_password(password),
                                int(approved),
                                now(),
                            ),
                        )
                        st.success("Doctor added.")
                        st.rerun()
                    except sqlite3.IntegrityError:
                        st.error("An account with this email already exists.")

    with directory_tab:
        search = st.text_input(
            "Search doctors",
            placeholder="Name, clinic, email, or phone",
            key="doctor_directory_search",
        ).strip()
        status_filter = st.selectbox(
            "Approval status",
            ["All", "Approved", "Awaiting approval"],
            key="doctor_status_filter",
        )
    doctors = query(
        "SELECT * FROM users WHERE role = 'doctor' ORDER BY approved, created_at DESC"
    )
    with directory_tab:
        filtered_doctors = []
        for doctor in doctors:
            haystack = " ".join(
                [
                    doctor["full_name"],
                    doctor["clinic_name"],
                    doctor["email"],
                    doctor["phone"],
                ]
            ).lower()
            status_matches = (
                status_filter == "All"
                or (status_filter == "Approved" and doctor["approved"])
                or (status_filter == "Awaiting approval" and not doctor["approved"])
            )
            if (not search or search.lower() in haystack) and status_matches:
                filtered_doctors.append(doctor)

        if not filtered_doctors:
            st.info("No doctors match the selected filters.")
        for doctor in filtered_doctors:
            with st.container(border=True):
                col1, col2, col3 = st.columns([3, 2, 1])
                col1.markdown(f"**{doctor['full_name']}**")
                col1.caption(
                    f"{doctor['clinic_name']} | {doctor['email']} | {doctor['phone']}"
                )
                col2.write("Approved" if doctor["approved"] else "Awaiting approval")
                if not doctor["approved"]:
                    if col3.button(
                        "Approve", key=f"approve_{doctor['id']}", type="primary"
                    ):
                        execute(
                            "UPDATE users SET approved = 1 WHERE id = ?", (doctor["id"],)
                        )
                        st.rerun()
                elif col3.button("Suspend", key=f"suspend_{doctor['id']}"):
                    execute(
                        "UPDATE users SET approved = 0 WHERE id = ?", (doctor["id"],)
                    )
                    st.rerun()


def manage_orders() -> None:
    st.subheader("Manage orders")
    add_tab, edit_tab = st.tabs(["Add order manually", "Edit existing order"])
    with add_tab:
        create_order_form()

    with edit_tab:
        selected_doctor = doctor_lookup(
            "Filter by doctor",
            "manage_order_doctor",
            allow_all=True,
        )
        order_id, _ = order_selector(
            doctor_id=selected_doctor,
            key="manage_order",
        )
        if order_id:
            order = query_one(
                """
                SELECT o.*, u.full_name AS doctor_name
                FROM orders o JOIN users u ON u.id = o.doctor_id WHERE o.id = ?
                """,
                (order_id,),
            )
            with st.container(border=True):
                st.write(
                    f"**{order['order_number']}** | {order['doctor_name']} | "
                    f"{order['patient_name']} | {order['service_type']}"
                )
                col1, col2 = st.columns(2)
                current_status = ORDER_STATUSES.index(order["status"])
                status = col1.selectbox(
                    "Order status", ORDER_STATUSES, index=current_status
                )
                amount = col2.number_input(
                    "Order amount (INR)",
                    min_value=0.0,
                    value=float(order["amount"]),
                    step=500.0,
                )
                notes = st.text_area("Notes", value=order["notes"])
                if st.button("Update order", type="primary"):
                    execute(
                        """
                        UPDATE orders
                        SET status = ?, amount = ?, notes = ?, updated_at = ?
                        WHERE id = ?
                        """,
                        (status, amount, notes, now(), order_id),
                    )
                    st.success("Order updated.")
                    st.rerun()

            uploads = query(
                "SELECT * FROM uploads WHERE order_id = ? ORDER BY uploaded_at DESC",
                (order_id,),
            )
            if uploads:
                st.markdown("**Uploaded files**")
                for item in uploads:
                    path = Path(item["stored_path"])
                    if path.exists():
                        st.download_button(
                            f"{item['category']}: {item['original_name']}",
                            data=path.read_bytes(),
                            file_name=item["original_name"],
                            key=f"download_upload_{item['id']}",
                        )


def manage_payments(user_id: int) -> None:
    st.subheader("Payments")
    selected_doctor = doctor_lookup(
        "Search doctor",
        "payment_doctor_lookup",
        allow_all=True,
    )
    order_id, _ = order_selector(
        doctor_id=selected_doctor,
        key="payment_order",
    )
    if not order_id:
        return
    order = query_one("SELECT * FROM orders WHERE id = ?", (order_id,))
    balance = max(float(order["amount"]) - float(order["amount_paid"]), 0)
    st.info(
        f"Order amount: {money(float(order['amount']))} | "
        f"Paid: {money(float(order['amount_paid']))} | Balance: {money(balance)}"
    )
    with st.form("record_payment", clear_on_submit=True):
        amount = st.number_input(
            "Payment amount", min_value=0.01, max_value=max(balance, 0.01), step=100.0
        )
        method = st.selectbox("Method", ["UPI", "Bank transfer", "Cash", "Card", "Other"])
        reference = st.text_input("Reference")
        notes = st.text_input("Notes")
        if st.form_submit_button("Record payment", type="primary"):
            new_paid = float(order["amount_paid"]) + amount
            if new_paid >= float(order["amount"]):
                payment_status = "Paid"
            elif new_paid > 0:
                payment_status = "Partial"
            else:
                payment_status = "Unpaid"
            execute(
                """
                INSERT INTO payments
                    (order_id, amount, method, reference, notes, paid_at, recorded_by)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (order_id, amount, method, reference, notes, now(), user_id),
            )
            execute(
                "UPDATE orders SET amount_paid = ?, payment_status = ?, updated_at = ? WHERE id = ?",
                (new_paid, payment_status, now(), order_id),
            )
            st.success("Payment recorded.")
            st.rerun()

    payments = query(
        "SELECT paid_at, amount, method, reference, notes FROM payments WHERE order_id = ? ORDER BY paid_at DESC",
        (order_id,),
    )
    if payments:
        st.dataframe(pd.DataFrame([dict(p) for p in payments]), hide_index=True, use_container_width=True)


def invoices_panel(user: sqlite3.Row) -> None:
    st.subheader("Invoices")
    if user["role"] == "admin":
        existing_tab, manual_tab, browse_tab = st.tabs(
            ["Generate from order", "Add invoice manually", "Invoice register"]
        )
    else:
        existing_tab, browse_tab = st.tabs(["Order invoice", "Invoice register"])
        manual_tab = None

    with existing_tab:
        selected_doctor = (
            doctor_lookup(
                "Search doctor",
                "invoice_doctor_lookup",
                allow_all=True,
            )
            if user["role"] == "admin"
            else user["id"]
        )
        order_id, _ = order_selector(
            selected_doctor,
            key="invoice_order",
            uninvoiced_only=user["role"] == "admin",
        )
        if order_id:
            invoice = query_one("SELECT * FROM invoices WHERE order_id = ?", (order_id,))
            if not invoice and user["role"] == "admin":
                col1, col2 = st.columns(2)
                tax_rate = col1.number_input(
                    "Tax rate (%)",
                    min_value=0.0,
                    max_value=100.0,
                    value=18.0,
                    key="existing_invoice_tax",
                )
                custom_number = col2.text_input(
                    "Invoice number (optional)",
                    placeholder="Auto-generated when blank",
                    key="existing_invoice_number",
                )
                if st.button("Generate invoice", type="primary"):
                    try:
                        invoice = generate_invoice_files(
                            order_id, tax_rate, custom_number or None
                        )
                        st.success(f"Invoice {invoice['invoice_number']} generated.")
                        st.rerun()
                    except sqlite3.IntegrityError:
                        st.error("That invoice number is already in use.")
            elif not invoice:
                st.info("The invoice has not been generated yet.")

    if manual_tab is not None:
        with manual_tab:
            doctors = doctor_options()
            if not doctors:
                st.info("Add and approve a doctor before creating an invoice.")
            else:
                with st.form("manual_invoice", clear_on_submit=True):
                    doctor_label = st.selectbox(
                        "Search and select doctor *",
                        doctors,
                        help="Type a doctor name, clinic, or email to search.",
                    )
                    col1, col2 = st.columns(2)
                    invoice_number = col1.text_input(
                        "Invoice number",
                        placeholder="Auto-generated when blank",
                    )
                    invoice_date = col2.date_input(
                        "Invoice date", value=datetime.now().date()
                    )
                    patient_name = col1.text_input("Patient name *")
                    patient_code = col2.text_input("Patient ID / code")
                    service_type = col1.text_input("Service / description *")
                    tooth_details = col2.text_input("Tooth / arch details")
                    subtotal = col1.number_input(
                        "Subtotal (INR)", min_value=0.0, step=500.0
                    )
                    tax_rate = col2.number_input(
                        "Tax rate (%)", min_value=0.0, max_value=100.0, value=18.0
                    )
                    notes = st.text_area("Notes")
                    submitted = st.form_submit_button(
                        "Add invoice", type="primary"
                    )
                    if submitted:
                        if not patient_name.strip() or not service_type.strip():
                            st.error("Patient name and service are required.")
                        else:
                            timestamp = now()
                            order_id = execute(
                                """
                                INSERT INTO orders
                                    (order_number, doctor_id, patient_name,
                                     patient_code, service_type, tooth_details,
                                     notes, status, amount, due_date, created_at,
                                     updated_at)
                                VALUES (?, ?, ?, ?, ?, ?, ?, 'Completed', ?, '', ?, ?)
                                """,
                                (
                                    make_order_number(),
                                    doctors[doctor_label],
                                    patient_name.strip(),
                                    patient_code.strip(),
                                    service_type.strip(),
                                    tooth_details.strip(),
                                    notes.strip(),
                                    subtotal,
                                    timestamp,
                                    timestamp,
                                ),
                            )
                            try:
                                invoice = generate_invoice_files(
                                    order_id,
                                    tax_rate,
                                    invoice_number or None,
                                    invoice_date,
                                )
                                st.success(
                                    f"Invoice {invoice['invoice_number']} added."
                                )
                                st.rerun()
                            except sqlite3.IntegrityError:
                                execute("DELETE FROM orders WHERE id = ?", (order_id,))
                                st.error("That invoice number is already in use.")

    with browse_tab:
        doctor_id = (
            doctor_lookup(
                "Filter invoices by doctor",
                "invoice_register_doctor",
                allow_all=True,
            )
            if user["role"] == "admin"
            else user["id"]
        )
        where = "WHERE o.doctor_id = ?" if doctor_id else ""
        params = (doctor_id,) if doctor_id else ()
        invoices = query(
            f"""
            SELECT i.*, o.order_number, o.patient_name, u.full_name AS doctor_name
            FROM invoices i
            JOIN orders o ON o.id = i.order_id
            JOIN users u ON u.id = o.doctor_id
            {where}
            ORDER BY i.created_at DESC
            """,
            params,
        )
        search = st.text_input(
            "Search invoice register",
            placeholder="Invoice, order, doctor, or patient",
            key="invoice_register_search",
        ).strip().lower()
        visible_invoices = [
            invoice
            for invoice in invoices
            if not search
            or search
            in " ".join(
                [
                    invoice["invoice_number"],
                    invoice["order_number"],
                    invoice["doctor_name"],
                    invoice["patient_name"],
                ]
            ).lower()
        ]
        if not visible_invoices:
            st.info("No invoices match the selected lookup.")
        for invoice in visible_invoices:
            with st.container(border=True):
                st.write(
                    f"**{invoice['invoice_number']}** | {invoice['doctor_name']} | "
                    f"{invoice['patient_name']} | {money(float(invoice['total']))}"
                )
                col1, col2 = st.columns(2)
                pdf_path = Path(invoice["pdf_path"])
                docx_path = Path(invoice["docx_path"])
                if pdf_path.exists():
                    col1.download_button(
                        "Download PDF",
                        pdf_path.read_bytes(),
                        file_name=pdf_path.name,
                        mime="application/pdf",
                        use_container_width=True,
                        key=f"invoice_pdf_{invoice['id']}",
                    )
                if docx_path.exists():
                    col2.download_button(
                        "Download DOCX",
                        docx_path.read_bytes(),
                        file_name=docx_path.name,
                        mime=(
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                        use_container_width=True,
                        key=f"invoice_docx_{invoice['id']}",
                    )


def render_app() -> None:
    st.set_page_config(page_title="Dental CRM", page_icon="D", layout="wide")
    st.markdown(
        """
        <style>
        :root { --ink: #123b5d; --aqua: #10a7a2; }
        .stApp { background: #f6f9fc; }
        .brand { font-size: 2.1rem; font-weight: 800; color: var(--ink); letter-spacing: -.04em; }
        div[data-testid="stMetric"] {
            background: white; border: 1px solid #e2e8f0; border-radius: 14px; padding: 16px;
        }
        div[data-testid="stForm"], div[data-testid="stVerticalBlockBorderWrapper"] {
            background: white; border-radius: 14px;
        }
        .stButton > button[kind="primary"], .stFormSubmitButton > button[kind="primary"] {
            background: var(--aqua); border-color: var(--aqua);
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    initialize_app()
    if "user_id" not in st.session_state:
        login_page()
        return

    user = query_one("SELECT * FROM users WHERE id = ?", (st.session_state.user_id,))
    if not user:
        del st.session_state.user_id
        st.rerun()

   from pathlib import Path

LOGO = Path(__file__).parent / "assets" / "vinayak-logo.jpg"

with st.sidebar:
    st.markdown(
        '<div class="brand">Dental CRM</div>',
        unsafe_allow_html=True
    )

    st.write(f"**{user['full_name']}**")
    st.caption(f"{user['role'].title()} | {user['clinic_name']}")

    if st.button("Logout", use_container_width=True):
        del st.session_state["user_id"]
        st.rerun()
        
        if user["role"] == "admin":
            pages = [
                "Dashboard",
                "Doctors",
                "Create order",
                "Manage orders",
                "Upload STL",
                "Upload Photos",
                "Invoices",
                "Payments",
            ]
        else:
            pages = [
                "Dashboard",
                "Create order",
                "Track orders",
                "Upload STL",
                "Upload Photos",
                "Invoices",
            ]
        page = st.radio("Navigation", pages, label_visibility="collapsed")

    st.title(page)
    if page == "Dashboard":
        admin_dashboard() if user["role"] == "admin" else doctor_dashboard(user["id"])
    elif page == "Doctors":
        manage_doctors()
    elif page == "Create order":
        create_order_form(None if user["role"] == "admin" else user["id"])
    elif page in ("Manage orders", "Track orders"):
        if user["role"] == "admin":
            manage_orders()
        else:
            frame = orders_table(user["id"])
            if frame.empty:
                st.info("No orders yet.")
            else:
                st.dataframe(frame, use_container_width=True, hide_index=True)
    elif page == "Upload STL":
        upload_panel(user, "STL")
    elif page == "Upload Photos":
        upload_panel(user, "Photo")
    elif page == "Invoices":
        invoices_panel(user)
    elif page == "Payments":
        manage_payments(user["id"])


if __name__ == "__main__":
    render_app()
